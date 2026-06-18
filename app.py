"""
Alpie Vision Tester — Streamlit app for testing the vision endpoint.
Supports image & video uploads via the OpenAI-compatible chat/completions API.
"""

import base64
import io
import json
import os
import re
import tempfile
import time

import cv2
import requests
import streamlit as st
import yt_dlp
from PIL import Image
from components.base64_uploader import base64_uploader, uploaded_files_to_images, uploaded_file_to_bytes


# ── Page config ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Alpie · Vision Tester",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ───────────────────────────────────────────────────────────────

st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

    :root {
        --bg-primary: #0a0a0f;
        --bg-card: rgba(18, 18, 28, 0.85);
        --bg-glass: rgba(255,255,255,0.04);
        --accent-1: #7c3aed;
        --accent-2: #06b6d4;
        --accent-3: #f472b6;
        --text-primary: #f1f5f9;
        --text-secondary: #94a3b8;
        --border: rgba(255,255,255,0.08);
        --glow: 0 0 30px rgba(124, 58, 237, .25);
    }

    html, body, [data-testid="stAppViewContainer"] {
        font-family: 'Inter', sans-serif !important;
        background: var(--bg-primary) !important;
        color: var(--text-primary) !important;
    }

    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0f0f1a 0%, #141425 100%) !important;
        border-right: 1px solid var(--border) !important;
    }

    /* Hero */
    .hero {
        text-align: center;
        padding: 3rem 1rem 2rem;
        background: radial-gradient(circle at center, rgba(124, 58, 237, 0.06) 0%, transparent 70%);
        margin-bottom: 1.5rem;
    }
    .hero h1 {
        font-size: 3.5rem;
        font-weight: 900;
        background: linear-gradient(135deg, var(--accent-1), var(--accent-2), var(--accent-3));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin: 0;
        letter-spacing: -1.5px;
        filter: drop-shadow(0 4px 12px rgba(124, 58, 237, 0.15));
    }
    .hero p {
        color: var(--text-secondary);
        font-size: 1.15rem;
        margin-top: .6rem;
        font-weight: 400;
        letter-spacing: 0.5px;
    }

    /* Section label */
    .section-label {
        font-weight: 700;
        font-size: .85rem;
        text-transform: uppercase;
        letter-spacing: 1.5px;
        color: var(--accent-2);
        margin-bottom: .6rem;
        display: flex;
        align-items: center;
        gap: 6px;
    }

    /* Think Details & Content */
    details.think-details {
        background: rgba(124, 58, 237, 0.04);
        border: 1px solid rgba(124, 58, 237, 0.15);
        border-radius: 12px;
        padding: 0.8rem 1.2rem;
        margin-bottom: 1.2rem;
        transition: all 0.3s ease;
    }
    details.think-details[open] {
        background: rgba(124, 58, 237, 0.08);
        border-color: rgba(124, 58, 237, 0.3);
        box-shadow: 0 4px 20px rgba(124, 58, 237, 0.1);
    }
    details.think-details summary {
        font-weight: 600;
        font-size: 0.9rem;
        color: #c4b5fd;
        cursor: pointer;
        user-select: none;
        outline: none;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    details.think-details summary::-webkit-details-marker {
        display: none;
    }
    details.think-details summary::before {
        content: "🔮";
        transition: transform 0.2s;
    }
    details.think-details[open] summary::before {
        transform: rotate(45deg);
    }
    .think-content {
        margin-top: 0.8rem;
        font-size: 0.92rem;
        line-height: 1.7;
        color: #d8b4fe;
        white-space: pre-wrap;
        border-left: 2px solid var(--accent-1);
        padding-left: 1rem;
    }

    /* Answer block */
    .answer-block {
        background: linear-gradient(180deg, rgba(255, 255, 255, 0.05) 0%, rgba(255, 255, 255, 0.02) 100%);
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-left: 4px solid var(--accent-2);
        border-radius: 12px;
        padding: 1.5rem;
        font-size: 1.05rem;
        line-height: 1.8;
        color: var(--text-primary);
        box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.37);
        backdrop-filter: blur(4px);
        -webkit-backdrop-filter: blur(4px);
        margin-bottom: 1.5rem;
    }

    /* Stat cards */
    .stat-row {
        display: flex;
        gap: 1rem;
        margin-top: 1.2rem;
        flex-wrap: wrap;
    }
    .stat-card {
        flex: 1;
        min-width: 140px;
        background: rgba(255, 255, 255, 0.03);
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 0.8rem 1rem;
        text-align: center;
        transition: transform 0.2s, border-color 0.2s;
    }
    .stat-card:hover {
        transform: translateY(-2px);
        border-color: rgba(6, 182, 212, 0.3);
    }
    .stat-card .val {
        font-size: 1.3rem;
        font-weight: 800;
        background: linear-gradient(90deg, var(--accent-2), var(--accent-1));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .stat-card .lbl {
        font-size: .75rem;
        color: var(--text-secondary);
        text-transform: uppercase;
        letter-spacing: 1.2px;
        margin-top: 4px;
        font-weight: 500;
    }

    /* Tabs */
    [data-testid="stTabs"] {
        background: rgba(18, 18, 28, 0.5);
        padding: 0.5rem 1rem 0 1rem;
        border-radius: 12px 12px 0 0;
        border: 1px solid var(--border);
        border-bottom: none;
    }
    [data-testid="stTabs"] button {
        font-family: 'Inter', sans-serif !important;
        font-weight: 600 !important;
        font-size: .95rem !important;
        color: var(--text-secondary) !important;
        border-bottom: 2px solid transparent !important;
        transition: all .2s ease !important;
        padding: 0.6rem 1.2rem !important;
    }
    [data-testid="stTabs"] button[aria-selected="true"] {
        color: var(--accent-2) !important;
        border-bottom: 2px solid var(--accent-2) !important;
    }

    /* Buttons */
    .stButton > button {
        background: linear-gradient(135deg, var(--accent-1), var(--accent-2)) !important;
        color: #fff !important;
        border: none !important;
        border-radius: 10px !important;
        font-weight: 700 !important;
        padding: 0.7rem 1.6rem !important;
        font-size: 0.95rem !important;
        transition: all .2s ease !important;
        box-shadow: 0 4px 15px rgba(124, 58, 237, 0.2) !important;
    }
    .stButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 22px rgba(124,58,237,.45) !important;
        filter: brightness(1.1);
    }
    .stButton > button:active {
        transform: translateY(0) !important;
    }

    /* Text input and area styling to fix prompt visibility in all modes and add gorgeous focus glow */
    [data-baseweb="textarea"], [data-baseweb="base-input"] {
        background-color: #12121e !important;
        background: #12121e !important;
        border-radius: 10px !important;
    }
    textarea, input, [data-baseweb="base-input"] input, [data-baseweb="textarea"] textarea {
        background-color: #12121e !important;
        background: #12121e !important;
        border: 1px solid var(--border) !important;
        border-radius: 10px !important;
        color: var(--text-primary) !important;
        -webkit-text-fill-color: var(--text-primary) !important;
        font-family: 'Inter', sans-serif !important;
        font-size: 0.95rem !important;
        transition: border-color 0.2s, box-shadow 0.2s !important;
    }
    textarea:focus, input:focus, [data-baseweb="base-input"] input:focus, [data-baseweb="textarea"] textarea:focus {
        border-color: var(--accent-2) !important;
        box-shadow: 0 0 10px rgba(6, 182, 212, 0.25) !important;
    }

    /* File uploader */
    [data-testid="stFileUploader"] {
        border: 2px dashed var(--border) !important;
        border-radius: 12px !important;
        padding: 1.5rem !important;
        background: rgba(255, 255, 255, 0.01) !important;
        transition: border-color .2s, background-color .2s !important;
    }
    [data-testid="stFileUploader"]:hover {
        border-color: var(--accent-1) !important;
        background-color: rgba(124, 58, 237, 0.02) !important;
    }

    /* Scrollbar */
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-track { background: transparent; }
    ::-webkit-scrollbar-thumb { background: rgba(124,58,237,.35); border-radius: 3px; }

    /* Image preview */
    .preview-grid {
        display: flex;
        gap: .6rem;
        flex-wrap: wrap;
        margin: .8rem 0;
    }
    .preview-grid img {
        border-radius: 10px;
        border: 1px solid var(--border);
        max-height: 200px;
        object-fit: cover;
    }

    /* Thinking / Typing Bubbles */
    .thinking-bubble, .direct-bubble {
        display: inline-flex;
        align-items: center;
        gap: 8px;
        border-radius: 20px;
        padding: 8px 16px;
        font-size: 0.9rem;
        font-weight: 600;
        animation: pulse-glow 2s infinite ease-in-out;
        margin-bottom: 1rem;
    }
    .thinking-bubble {
        background: rgba(124, 58, 237, 0.08) !important;
        border: 1px solid rgba(124, 58, 237, 0.25) !important;
        color: #c4b5fd !important;
        box-shadow: 0 0 15px rgba(124, 58, 237, 0.15);
    }
    .direct-bubble {
        background: rgba(6, 182, 212, 0.08) !important;
        border: 1px solid rgba(6, 182, 212, 0.25) !important;
        color: #a5f3fc !important;
        box-shadow: 0 0 15px rgba(6, 182, 212, 0.15);
    }
    .thinking-bubble .dots, .direct-bubble .dots {
        display: inline-flex;
        gap: 4px;
        align-items: center;
    }
    .thinking-bubble .dot, .direct-bubble .dot {
        width: 6px;
        height: 6px;
        border-radius: 50%;
        animation: dot-bounce 1.4s infinite ease-in-out both;
    }
    .thinking-bubble .dot {
        background-color: #c4b5fd !important;
    }
    .direct-bubble .dot {
        background-color: #a5f3fc !important;
    }
    .thinking-bubble .dot:nth-child(1), .direct-bubble .dot:nth-child(1) { animation-delay: -0.32s; }
    .thinking-bubble .dot:nth-child(2), .direct-bubble .dot:nth-child(2) { animation-delay: -0.16s; }

    @keyframes pulse-glow {
        0%, 100% { opacity: 0.8; transform: translateY(0); }
        50% { opacity: 1; transform: translateY(-2px); }
    }
    @keyframes dot-bounce {
        0%, 80%, 100% { transform: scale(0.3); opacity: 0.3; }
        40% { transform: scale(1.1); opacity: 1; }
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ── Constants ────────────────────────────────────────────────────────────────

API_URL = os.environ.get("API_URL", "http://20.245.200.125:8000/v1/chat/completions")
MAX_FRAMES = 8

# Check for system-wide/default cookies at startup
if "default_cookies_path" not in st.session_state:
    default_path = None
    # 1. Check if raw cookies content is in environment variables (useful for secrets)
    raw_cookies = os.environ.get("YOUTUBE_COOKIES")
    if raw_cookies:
        try:
            raw_cookies_str = raw_cookies.strip()
            # Handle potential base64 encoded cookies to prevent format loss in key-value editors
            if not raw_cookies_str.startswith("#") and len(raw_cookies_str) % 4 == 0:
                try:
                    raw_cookies_str = base64.b64decode(raw_cookies_str).decode("utf-8").strip()
                except Exception:
                    pass
            
            # Normalize all newlines to Unix format \n for the Linux container
            raw_cookies_str = raw_cookies_str.replace("\r\n", "\n").replace("\r", "\n")
            
            temp_cookies_path = os.path.join(tempfile.gettempdir(), "system_yt_cookies.txt")
            with open(temp_cookies_path, "w", newline="\n", encoding="utf-8") as f:
                f.write(raw_cookies_str)
            
            # Diagnostic logs for container output
            file_size = os.path.getsize(temp_cookies_path)
            print(f"[debug] YOUTUBE_COOKIES env loaded. Wrote to: {temp_cookies_path} ({file_size} bytes)")
            if file_size > 0:
                with open(temp_cookies_path, "r", encoding="utf-8") as rf:
                    head = rf.read(100)
                print(f"[debug] Cookies file header signature: {repr(head)}")
            
            default_path = temp_cookies_path
        except Exception as e:
            st.sidebar.error(f"Error loading YOUTUBE_COOKIES env: {e}")
            print(f"[debug] Error loading YOUTUBE_COOKIES env: {e}")
            
    # 2. Check for local cookies.txt file in app directory or /root
    if not default_path:
        for p in ["cookies.txt", "cookie.txt", "/root/cookies.txt", "/root/cookie.txt"]:
            if os.path.exists(p):
                default_path = os.path.abspath(p)
                file_size = os.path.getsize(default_path)
                print(f"[debug] Found local cookies file: {default_path} ({file_size} bytes)")
                break
                
    st.session_state["default_cookies_path"] = default_path

# Initialize session state for thinking toggle
if "enable_thinking" not in st.session_state:
    st.session_state.enable_thinking = True

# Auto-detect model name from the server
if "model_name" not in st.session_state:
    st.session_state.model_name = ""
    # Try to auto-fetch on first load
    try:
        _r = requests.get(
            API_URL.replace("/v1/chat/completions", "/v1/models"),
            timeout=5,
        )
        if _r.status_code == 200:
            _data = _r.json()
            _models = _data.get("data", [])
            if _models:
                st.session_state.model_name = _models[0].get("id", "")
    except Exception:
        pass

# ── Helpers ──────────────────────────────────────────────────────────────────


def image_to_base64_url(img: Image.Image, fmt: str = "JPEG", max_side: int | None = 1280) -> str:
    """Convert a PIL Image to a base64 data-URL string.

    To avoid AxiosError / payload-too-large errors, images are:
    1. Converted to RGB (drops alpha channel so JPEG works on PNGs with transparency)
    2. Down-scaled so the longest side is at most *max_side* pixels (if max_side is not None)
    3. Encoded as JPEG at quality 85 (much smaller than PNG)
    """
    # Ensure RGB — JPEG cannot encode RGBA / palette modes
    if img.mode not in ("RGB",):
        img = img.convert("RGB")

    # Resize if either dimension exceeds max_side
    if max_side is not None:
        w, h = img.size
        if max(w, h) > max_side:
            scale = max_side / max(w, h)
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    b64 = base64.b64encode(buf.getvalue()).decode()
    return f"data:image/jpeg;base64,{b64}"


def extract_video_id(url: str) -> str | None:
    """Extract YouTube video ID from various URL formats."""
    patterns = [
        r'(?:v=|/v/|youtu\.be/)([a-zA-Z0-9_-]{11})',
        r'(?:embed/)([a-zA-Z0-9_-]{11})',
        r'(?:shorts/)([a-zA-Z0-9_-]{11})',
    ]
    for pattern in patterns:
        m = re.search(pattern, url)
        if m:
            return m.group(1)
    return None


def extract_youtube_thumbnails(url: str) -> list[Image.Image]:
    """Fetch thumbnail frames from YouTube's public CDN.

    YouTube serves multiple thumbnail images at predictable URLs
    that are accessible from ANY IP (including datacenter) without
    authentication.  This is the ultimate fallback when yt-dlp is
    blocked by bot detection.
    """
    video_id = extract_video_id(url)
    if not video_id:
        return []

    # Thumbnail URLs ordered from best to worst quality
    thumb_urls = [
        f"https://img.youtube.com/vi/{video_id}/maxresdefault.jpg",
        f"https://img.youtube.com/vi/{video_id}/sddefault.jpg",
        f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg",
        f"https://img.youtube.com/vi/{video_id}/mqdefault.jpg",
        f"https://img.youtube.com/vi/{video_id}/0.jpg",   # default poster
        f"https://img.youtube.com/vi/{video_id}/1.jpg",   # ~25% mark
        f"https://img.youtube.com/vi/{video_id}/2.jpg",   # ~50% mark
        f"https://img.youtube.com/vi/{video_id}/3.jpg",   # ~75% mark
    ]

    frames: list[Image.Image] = []
    seen_sizes: set[tuple[int, int]] = set()  # deduplicate identical placeholders
    for thumb_url in thumb_urls:
        try:
            r = requests.get(thumb_url, timeout=10)
            if r.status_code == 200 and len(r.content) > 1000:  # skip tiny placeholder imgs
                img = Image.open(io.BytesIO(r.content)).convert("RGB")
                # YouTube returns a grey 120x90 placeholder for missing thumbnails
                if img.size[0] >= 200 and img.size not in seen_sizes:
                    frames.append(img)
                    seen_sizes.add(img.size)
        except Exception:
            continue
    return frames


def download_youtube_video(url: str, quality: str = "Lowest (Fastest)") -> tuple[str | None, list[Image.Image] | None]:
    """Download lowest or selected resolution stream of a YouTube or generic video to a temp file.

    Returns (filepath, None) on success, (None, thumbnail_frames) if only
    thumbnails could be fetched, or (None, None) on total failure.

    Uses a multi-strategy approach to bypass YouTube bot detection:
    1. User-provided cookies.txt file (most reliable)
    2. Browser cookies via cookiesfrombrowser (local dev only)
    3. Cookieless download with various player clients
    4. Fallback: YouTube thumbnail CDN (always works)
    """
    temp_dir = tempfile.gettempdir()
    print(f"[debug] download_youtube_video called for URL: {url}")

    # Strategy 1: User-provided or system-wide cookies file (most reliable)
    cookies_path = st.session_state.get("yt_cookies_path") or st.session_state.get("default_cookies_path")
    print(f"[debug] Resolved cookies_path: {cookies_path}")
    if cookies_path:
        if os.path.exists(cookies_path):
            file_size = os.path.getsize(cookies_path)
            print(f"[debug] Cookies file exists: {cookies_path} ({file_size} bytes)")
            try:
                with open(cookies_path, "r", encoding="utf-8") as f:
                    print(f"[debug] First line of cookies file: {repr(f.readline())}")
            except Exception as ce:
                print(f"[debug] Failed to read cookies file line: {ce}")
        else:
            print(f"[debug] Cookies path specified but file does not exist: {cookies_path}")

    # Try importing ImpersonateTarget for browser impersonation (requires curl_cffi)
    impersonate_target = None
    try:
        import curl_cffi
        from yt_dlp.networking.impersonate import ImpersonateTarget
        impersonate_target = ImpersonateTarget.from_str("chrome-116:windows-10")
    except ImportError:
        pass
    except Exception:
        pass

    format_queries = {
        "Lowest (Fastest)": "worst[ext=mp4]/lowest[ext=mp4]/worst/lowest",
        "Medium (Balanced)": "bestvideo[height<=720][ext=mp4]+bestaudio[ext=m4a]/best[height<=720]/worst/lowest",
        "Highest (Best Quality)": "bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best"
    }
    fmt = format_queries.get(quality, "worst[ext=mp4]/lowest[ext=mp4]/worst/lowest")

    base_opts = {
        'format': fmt,
        'outtmpl': 'yt_download_%(id)s.%(ext)s',
        'paths': {
            'home': temp_dir,
            'temp': temp_dir,
        },
        'nopart': True,
        'overwrites': True,
        'quiet': True,
        'no_warnings': True,
    }

    if impersonate_target:
        base_opts['impersonate'] = impersonate_target

    # Build list of strategies to try in order
    strategies = []

    # Strategy 1: User-provided or system-wide cookies file (most reliable)
    cookies_path = st.session_state.get("yt_cookies_path") or st.session_state.get("default_cookies_path")
    if cookies_path and os.path.exists(cookies_path):
        opts_cookies = {**base_opts, 'cookiefile': cookies_path}
        strategies.append(("cookies file", opts_cookies))

        opts_cookies_no_sdkless = {
            **base_opts,
            'cookiefile': cookies_path,
            'extractor_args': {'youtube': {'player_client': ['default', '-android_sdkless']}}
        }
        strategies.append(("cookies file (no sdkless)", opts_cookies_no_sdkless))

        opts_cookies_web = {
            **base_opts,
            'cookiefile': cookies_path,
            'extractor_args': {'youtube': {'player_client': ['web', 'mweb', 'android']}}
        }
        strategies.append(("cookies file (web/android)", opts_cookies_web))

    # Strategy 2: Browser cookies (works for local development)
    for browser in ['chrome', 'firefox', 'edge', 'brave']:
        opts_browser = {**base_opts, 'cookiesfrombrowser': (browser,)}
        strategies.append((f"{browser} browser cookies", opts_browser))

    # Strategy 3: Cookieless (default)
    strategies.append(("cookieless (default)", base_opts))

    # Strategy 3.5: Cookieless (excluding sdkless)
    opts_no_sdkless = {
        **base_opts,
        'extractor_args': {'youtube': {'player_client': ['default', '-android_sdkless']}}
    }
    strategies.append(("cookieless (no sdkless)", opts_no_sdkless))

    # Strategy 4: Cookieless (forcing iOS client)
    opts_ios = {
        **base_opts,
        'extractor_args': {'youtube': {'player_client': ['ios']}}
    }
    strategies.append(("cookieless (ios client)", opts_ios))

    # Strategy 4.5: Cookieless (forcing Android VR client)
    opts_vr = {
        **base_opts,
        'extractor_args': {'youtube': {'player_client': ['android_vr']}}
    }
    strategies.append(("cookieless (android_vr client)", opts_vr))

    # Strategy 4.6: Cookieless (forcing Web Safari client)
    opts_safari = {
        **base_opts,
        'extractor_args': {'youtube': {'player_client': ['web_safari']}}
    }
    strategies.append(("cookieless (web_safari client)", opts_safari))

    # Strategy 4.7: Cookieless (forcing Web Embedded client)
    opts_embedded = {
        **base_opts,
        'extractor_args': {'youtube': {'player_client': ['web_embedded']}}
    }
    strategies.append(("cookieless (web_embedded client)", opts_embedded))

    # Strategy 5: Cookieless (forcing TV client)
    opts_tv = {
        **base_opts,
        'extractor_args': {'youtube': {'player_client': ['tv']}}
    }
    strategies.append(("cookieless (tv client)", opts_tv))

    last_error = None
    for strategy_name, ydl_opts in strategies:
        try:
            print(f"[debug] Attempting strategy: '{strategy_name}'...")
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                # Check if already downloaded
                info = ydl.extract_info(url, download=False)
                filepath = ydl.prepare_filename(info)
                if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
                    print(f"[debug] Video already downloaded: {filepath}")
                    return filepath, None

                # Download from scratch
                info = ydl.extract_info(url, download=True)
                filepath = ydl.prepare_filename(info)
                if os.path.exists(filepath):
                    print(f"[debug] Download succeeded under '{strategy_name}': {filepath}")
                    return filepath, None
                base_path = os.path.splitext(filepath)[0]
                for ext in ['mp4', 'mkv', 'webm', '3gp', 'flv']:
                    p = f"{base_path}.{ext}"
                    if os.path.exists(p):
                        print(f"[debug] Download succeeded under '{strategy_name}' (as format): {p}")
                        return p, None
        except Exception as e:
            last_error = e
            print(f"[debug] Strategy '{strategy_name}' failed: {e}")
            continue  # Try next strategy

    # All yt-dlp strategies failed — fall back to YouTube thumbnail CDN
    st.warning(
        "⚠️ Video download blocked by YouTube bot detection (common on cloud servers). "
        "Falling back to thumbnail frames. For full video analysis, upload a `cookies.txt` "
        "file in the sidebar."
    )
    thumb_frames = extract_youtube_thumbnails(url)
    if thumb_frames:
        return None, thumb_frames

    # Total failure
    st.error(f"❌ Failed to download YouTube video: {last_error}")
    st.info(
        "💡 **Tip:** To bypass YouTube bot detection, export your browser cookies as a "
        "`cookies.txt` file (Netscape format) and upload it in the sidebar under **🎬 Video**. "
        "You can use a browser extension like \"Get cookies.txt LOCALLY\" to export them."
    )
    return None, None


def extract_frames(path: str, n: int = MAX_FRAMES) -> list[Image.Image]:
    """Extract *n* evenly-spaced frames from a video file."""
    cap = cv2.VideoCapture(path)
    total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    if total <= 0:
        cap.release()
        return []
    idxs = [int(i * total / n) for i in range(n)]
    frames: list[Image.Image] = []
    for idx in idxs:
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ok, frame = cap.read()
        if ok:
            frames.append(Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)))
    cap.release()
    return frames


def parse_output(raw: str) -> tuple[str | None, str]:
    """Split <think>…</think> from the final answer, supporting partial streaming states."""
    if not raw:
        return None, ""
    raw = raw.replace("<|im_end|>", "").replace("<|endoftext|>", "").strip()
    
    # Case 1: We have both tags
    if "<think>" in raw and "</think>" in raw:
        parts = raw.split("<think>", 1)[1].split("</think>", 1)
        think = parts[0].strip()
        answer = parts[1].strip()
        return think or None, answer
        
    # Case 2: We have <think> but it is not closed yet (streaming)
    if "<think>" in raw:
        think = raw.split("<think>", 1)[1].strip()
        return think or None, ""
        
    # Case 3: No think tags at all
    return None, raw


def call_api(
    messages: list[dict],
    model: str = "",
    max_tokens: int = 8192,
    temperature: float = 0.7,
    stream: bool = True,
    enable_thinking: bool = True,
) -> requests.Response | dict:
    """
    Call the OpenAI-compatible chat/completions endpoint.
    Returns a streaming Response when stream=True, or parsed JSON otherwise.
    """
    payload = {
        "model": model,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": max(temperature, 0.01),
        "stream": stream,
        "chat_template_kwargs": {
            "enable_thinking": enable_thinking
        }
    }
    headers = {"Content-Type": "application/json"}

    if stream:
        resp = requests.post(API_URL, json=payload, headers=headers, stream=True, timeout=300)
        resp.raise_for_status()
        return resp
    else:
        resp = requests.post(API_URL, json=payload, headers=headers, timeout=300)
        resp.raise_for_status()
        return resp.json()


def stream_response(resp: requests.Response):
    """Yield (reasoning_text, answer_text, used_reasoning_field) from an SSE stream.

    vLLM sends thinking tokens in ``delta.reasoning`` (or ``delta.reasoning_content``
    depending on version) and final answer tokens in ``delta.content``.
    Older setups embed ``<think>`` tags inside ``delta.content`` instead.
    This function handles all variants.
    """
    reasoning_text = ""
    answer_text = ""
    used_reasoning_field = False  # track whether the API uses the dedicated field
    for line in resp.iter_lines(decode_unicode=True):
        if not line:
            continue
        if line.startswith("data: "):
            data_str = line[6:]
            if data_str.strip() == "[DONE]":
                break
            try:
                chunk = json.loads(data_str)
                delta = chunk.get("choices", [{}])[0].get("delta", {})
                # Capture reasoning / reasoning_content (vLLM reasoning-parser path)
                r_token = delta.get("reasoning") or delta.get("reasoning_content") or ""
                if r_token:
                    reasoning_text += r_token
                    used_reasoning_field = True
                # Capture regular content (can be None)
                c_token = delta.get("content") or ""
                if c_token:
                    answer_text += c_token
                if r_token or c_token:
                    yield reasoning_text, answer_text, used_reasoning_field
            except json.JSONDecodeError:
                continue
    yield reasoning_text, answer_text, used_reasoning_field


def build_image_messages(prompt: str, images: list[Image.Image], max_side: int | None = 1280) -> list[dict]:
    """Build a multimodal user message with text + images."""
    content: list[dict] = []
    for img in images:
        content.append({
            "type": "image_url",
            "image_url": {"url": image_to_base64_url(img, max_side=max_side)},
        })
    content.append({"type": "text", "text": prompt})
    return [{"role": "user", "content": content}]


def build_video_messages(prompt: str, frames: list[Image.Image], max_side: int | None = 1280) -> list[dict]:
    """Build a multimodal message from extracted video frames."""
    content: list[dict] = []
    for i, frame in enumerate(frames):
        content.append({
            "type": "image_url",
            "image_url": {"url": image_to_base64_url(frame, max_side=max_side)},
        })
    content.append({
        "type": "text",
        "text": f"[These are {len(frames)} frames extracted from a video]\n\n{prompt}",
    })
    return [{"role": "user", "content": content}]


def extract_text_from_file(file_info: dict) -> tuple[str, int]:
    """Extract text from base64 uploaded file dict.
    Returns (extracted_text, page_count).
    """
    import base64
    import io

    name = file_info.get("name", "").lower()
    b64_str = file_info.get("data", "")
    if "," in b64_str:
        b64_str = b64_str.split(",", 1)[1]

    try:
        file_bytes = base64.b64decode(b64_str)
    except Exception as e:
        return f"Error decoding base64 data for file {file_info.get('name')}: {e}", 0

    text = ""
    pages = 1

    try:
        if name.endswith(".pdf"):
            import pypdf
            pdf_file = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(pdf_file)
            pages = len(reader.pages)
            text_parts = []
            for i, page in enumerate(reader.pages):
                t = page.extract_text()
                if t:
                    text_parts.append(t)
            text = "\n\n".join(text_parts)

        elif name.endswith(".docx"):
            import docx
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)

            p_texts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    p_texts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    # Deduplicate cells since merged cells appear multiple times in row.cells
                    # but keep order as much as possible
                    seen = set()
                    unique_row = []
                    for val in row_text:
                        val_strip = val.strip()
                        if val_strip and val_strip not in seen:
                            seen.add(val_strip)
                            unique_row.append(val_strip)
                    if unique_row:
                        p_texts.append(" | ".join(unique_row))
            text = "\n\n".join(p_texts)

            # Approximate page count as words / 300
            words = len(text.split())
            pages = max(1, words // 300)

        else:
            # Fallback for plain text files (.txt) or other readable text files
            try:
                text = file_bytes.decode("utf-8", errors="ignore")
            except Exception:
                text = file_bytes.decode("latin1", errors="ignore")
            pages = max(1, len(text.split()) // 300)

    except Exception as e:
        return f"Error reading document content for {file_info.get('name')}: {str(e)}", 0

    return text.strip(), pages


def build_document_messages(prompt: str, documents: list[dict]) -> list[dict]:
    """Build a chat user message containing the extracted text of the documents and the prompt."""
    content_parts = []
    for doc in documents:
        content_parts.append(
            f"--- START OF DOCUMENT: {doc['name']} ---\n"
            f"{doc['text']}\n"
            f"--- END OF DOCUMENT: {doc['name']} ---\n"
        )
    content_parts.append(prompt)
    full_prompt = "\n\n".join(content_parts)
    return [{"role": "user", "content": full_prompt}]



def render_stream(messages: list[dict], max_tokens: int, temperature: float, model: str = "", enable_thinking: bool = True):
    """Call the API with streaming and render results live."""
    if not enable_thinking:
        system_instruction = {
            "role": "system",
            "content": "You are a direct assistant. Answer the user prompt directly without any thinking process. Do not write any thoughts inside <think> and </think> tags. Start answering directly."
        }
        messages = [system_instruction] + list(messages)

    indicator_ph = st.empty()
    think_label_ph = st.empty()
    think_ph = st.empty()
    ans_label_ph = st.empty()
    ans_ph = st.empty()
    status = st.empty()

    if enable_thinking:
        indicator_ph.markdown(
            """
            <div class="thinking-bubble">
                <span>🔮 Alpie is thinking</span>
                <div class="dots">
                    <div class="dot"></div>
                    <div class="dot"></div>
                    <div class="dot"></div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        indicator_ph.markdown(
            """
            <div class="direct-bubble">
                <span>⚡ Alpie is typing</span>
                <div class="dots">
                    <div class="dot"></div>
                    <div class="dot"></div>
                    <div class="dot"></div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    t0 = time.time()

    try:
        resp = call_api(messages, model=model, max_tokens=max_tokens, temperature=temperature, stream=True, enable_thinking=enable_thinking)
        final_reasoning = ""
        final_answer = ""
        token_count = 0

        for reasoning_text, answer_text, used_reasoning_field in stream_response(resp):
            final_reasoning = reasoning_text
            final_answer = answer_text
            token_count += 1

            if used_reasoning_field:
                # API provides reasoning_content separately (vLLM reasoning-parser)
                thinking = reasoning_text.strip() or None
                answer = answer_text.strip()
            else:
                # Fallback: parse <think> tags from the combined content
                thinking, answer = parse_output(answer_text)

            if thinking:
                think_label_ph.empty()
                think_ph.markdown(
                    f'<details open class="think-details"><summary>Chain of Thought</summary><div class="think-content">{thinking}</div></details>',
                    unsafe_allow_html=True,
                )

            if answer:
                indicator_ph.empty()
                ans_label_ph.markdown(
                    '<div class="section-label">✦ Answer</div>',
                    unsafe_allow_html=True,
                )
                ans_ph.markdown(
                    f'<div class="answer-block">{answer}</div>',
                    unsafe_allow_html=True,
                )

        elapsed = time.time() - t0
        total_chars = len(final_reasoning) + len(final_answer)
        status.empty()
        indicator_ph.empty()

        # Stats bar
        st.markdown(
            f"""
            <div class="stat-row">
                <div class="stat-card">
                    <div class="val">{elapsed:.1f}s</div>
                    <div class="lbl">Latency</div>
                </div>
                <div class="stat-card">
                    <div class="val">{total_chars}</div>
                    <div class="lbl">Characters</div>
                </div>
                <div class="stat-card">
                    <div class="val">{token_count}</div>
                    <div class="lbl">Chunks</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    except requests.exceptions.ConnectionError:
        status.empty()
        indicator_ph.empty()
        st.error("❌ Cannot connect to the API server. Is it running?")
    except requests.exceptions.Timeout:
        status.empty()
        indicator_ph.empty()
        st.error("❌ Request timed out after 300 seconds.")
    except requests.exceptions.HTTPError as e:
        status.empty()
        indicator_ph.empty()
        st.error(f"❌ HTTP Error: {e.response.status_code} — {e.response.text[:500]}")
    except Exception as e:
        status.empty()
        indicator_ph.empty()
        st.error(f"❌ {type(e).__name__}: {e}")


def render_non_stream(messages: list[dict], max_tokens: int, temperature: float, model: str = "", enable_thinking: bool = True):
    """Call the API without streaming and render the full result."""
    if not enable_thinking:
        system_instruction = {
            "role": "system",
            "content": "You are a direct assistant. Answer the user prompt directly without any thinking process. Do not write any thoughts inside <think> and </think> tags. Start answering directly."
        }
        messages = [system_instruction] + list(messages)

    status = st.empty()
    if enable_thinking:
        status.markdown(
            """
            <div class="thinking-bubble">
                <span>🔮 Alpie is thinking</span>
                <div class="dots">
                    <div class="dot"></div>
                    <div class="dot"></div>
                    <div class="dot"></div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        status.markdown(
            """
            <div class="direct-bubble">
                <span>⚡ Alpie is typing</span>
                <div class="dots">
                    <div class="dot"></div>
                    <div class="dot"></div>
                    <div class="dot"></div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    t0 = time.time()

    try:
        result = call_api(messages, model=model, max_tokens=max_tokens, temperature=temperature, stream=False, enable_thinking=enable_thinking)
        elapsed = time.time() - t0
        status.empty()

        message = result.get("choices", [{}])[0].get("message", {})
        raw = message.get("content") or ""
        reasoning_raw = message.get("reasoning") or message.get("reasoning_content") or ""
        usage = result.get("usage", {})

        if reasoning_raw.strip():
            # vLLM reasoning-parser provides reasoning separately
            thinking = reasoning_raw.strip()
            answer = raw.strip()
        else:
            # Fallback: parse <think> tags from content
            thinking, answer = parse_output(raw)

        if thinking:
            st.markdown(
                f'<details class="think-details"><summary>Chain of Thought</summary><div class="think-content">{thinking}</div></details>',
                unsafe_allow_html=True,
            )

        st.markdown(
            '<div class="section-label">✦ Answer</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            f'<div class="answer-block">{answer}</div>',
            unsafe_allow_html=True,
        )

        # Stats
        prompt_tok = usage.get("prompt_tokens", "—")
        compl_tok = usage.get("completion_tokens", "—")
        st.markdown(
            f"""
            <div class="stat-row">
                <div class="stat-card">
                    <div class="val">{elapsed:.1f}s</div>
                    <div class="lbl">Latency</div>
                </div>
                <div class="stat-card">
                    <div class="val">{prompt_tok}</div>
                    <div class="lbl">Prompt Tokens</div>
                </div>
                <div class="stat-card">
                    <div class="val">{compl_tok}</div>
                    <div class="lbl">Completion Tokens</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    except requests.exceptions.ConnectionError:
        status.empty()
        st.error("❌ Cannot connect to the API server. Is it running?")
    except requests.exceptions.Timeout:
        status.empty()
        st.error("❌ Request timed out after 300 seconds.")
    except requests.exceptions.HTTPError as e:
        status.empty()
        st.error(f"❌ HTTP Error: {e.response.status_code} — {e.response.text[:500]}")
    except Exception as e:
        status.empty()
        st.error(f"❌ {type(e).__name__}: {e}")


# ── Hero ─────────────────────────────────────────────────────────────────────

st.markdown(
    '<div class="hero">'
    "<h1>🧠 ALPIE</h1>"
    "<p>Vision Tester · Image · Video · Multimodal Reasoning &bull; 64k Context</p>"
    "</div>",
    unsafe_allow_html=True,
)

# ── Sidebar ──────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown(
        '<div class="section-label">🔗 Endpoint</div>',
        unsafe_allow_html=True,
    )
    st.code(API_URL, language=None)

    st.markdown("---")

    st.markdown(
        '<div class="section-label">🤖 Model</div>',
        unsafe_allow_html=True,
    )
    model_name = st.text_input(
        "Model name",
        value=st.session_state.model_name,
        placeholder="e.g. Qwen/Qwen2.5-VL-7B-Instruct",
        label_visibility="collapsed",
        help="The model ID served by vLLM. Click 'Check API Health' to auto-detect.",
    )
    st.session_state.model_name = model_name

    st.markdown("---")

    st.markdown(
        '<div class="section-label">⚙ Settings</div>',
        unsafe_allow_html=True,
    )

    enable_thinking = st.toggle(
        "🔮 Enable Thinking (Reasoning)",
        value=st.session_state.enable_thinking,
        help="Toggle reasoning/thinking mode (uses vLLM template configuration and direct prompt injection)",
    )
    st.session_state.enable_thinking = enable_thinking

    use_streaming = st.toggle(
        "🌊 Stream response",
        value=True,
        help="Stream tokens as they are generated",
    )

    st.markdown("---")

    st.markdown(
        '<div class="section-label">🎛 Generation</div>',
        unsafe_allow_html=True,
    )

    max_tokens = st.slider("Max New Tokens", 256, 65536, 8192, 256)
    temperature = st.slider("Temperature", 0.1, 1.5, 0.7, 0.05)

    st.markdown("---")

    st.markdown(
        '<div class="section-label">🖼 Image Settings</div>',
        unsafe_allow_html=True,
    )
    img_res_option = st.selectbox(
        "Max Image Dimension",
        ["Original", "2048", "1600", "1280 (Default)", "1024", "768", "512"],
        index=3,
        help="Scale down uploaded images or video frames. Higher resolutions capture more detail but increase payload size and latency."
    )
    if img_res_option == "Original":
        img_max_side = None
    else:
        img_max_side = int(img_res_option.split()[0])

    st.markdown("---")

    st.markdown(
        '<div class="section-label">🎬 Video Settings</div>',
        unsafe_allow_html=True,
    )
    max_frames = st.slider("Max Frames to Extract", 1, 16, MAX_FRAMES, 1)
    video_quality = st.selectbox(
        "Video Download Quality",
        ["Lowest (Fastest)", "Medium (Balanced)", "Highest (Best Quality)"],
        index=0,
        help="Select video resolution quality to download from URLs. Higher quality takes longer to download."
    )

    # Display status of system-wide cookies
    default_cookies_path = st.session_state.get("default_cookies_path")
    if default_cookies_path:
        if "system_yt_cookies" in default_cookies_path:
            st.info("🍪 **System cookies active** (loaded from environment variable).")
        else:
            st.info(f"🍪 **System cookies active** (loaded from `{os.path.basename(default_cookies_path)}`).")
    else:
        st.warning("⚠️ **No system cookies active.** Paste/upload below if blocked.")

    yt_cookies_data = base64_uploader(
        "🍪 YouTube cookies.txt (optional override)",
        accept=["txt"],
        multiple=False,
        key="yt_cookies",
    )
    if yt_cookies_data is not None:
        # Save the uploaded cookies to a temp file for yt-dlp
        cookie_file = yt_cookies_data[0] if isinstance(yt_cookies_data, list) else yt_cookies_data
        cookies_bytes = uploaded_file_to_bytes(cookie_file)
        cookies_path = os.path.join(tempfile.gettempdir(), "yt_cookies.txt")
        with open(cookies_path, "wb") as f:
            f.write(cookies_bytes)
        st.session_state["yt_cookies_path"] = cookies_path
    elif "yt_cookies_path" not in st.session_state:
        st.session_state["yt_cookies_path"] = None

    st.markdown("---")

    # Health check
    if st.button("🩺 Check API Health", use_container_width=True):
        try:
            r = requests.get(
                API_URL.replace("/v1/chat/completions", "/v1/models"),
                timeout=10,
            )
            if r.status_code == 200:
                models_data = r.json()
                st.success("✅ API is online!")
                st.json(models_data, expanded=False)
                # Auto-fill model name if empty
                available = models_data.get("data", [])
                if available and not model_name.strip():
                    st.session_state.model_name = available[0].get("id", "")
                    st.rerun()
            else:
                st.warning(f"⚠ Status {r.status_code}")
        except Exception as e:
            st.error(f"❌ {e}")

    st.markdown("---")
    xsrf_status = st.get_option("server.enableXsrfProtection")
    cors_status = st.get_option("server.enableCORS")
    st.markdown(
        f'<p style="font-size: 0.75rem; color: var(--text-secondary); text-align: center; margin-top: 10px;">'
        f'XSRF Protection: <strong style="color: {"#ef4444" if xsrf_status else "#10b981"}">{"Enabled" if xsrf_status else "Disabled"}</strong> &bull; '
        f'CORS Control: <strong style="color: {"#ef4444" if cors_status else "#10b981"}">{"Enabled" if cors_status else "Disabled"}</strong></p>',
        unsafe_allow_html=True
    )

# ── Status Bar ────────────────────────────────────────────────────────────────

enable_thinking = st.session_state.enable_thinking
active_mode = "🔮 Thinking Mode" if enable_thinking else "⚡ Direct Mode"
active_mode_color = "var(--accent-1)" if enable_thinking else "var(--accent-2)"
active_mode_bg = "rgba(124, 58, 237, 0.15)" if enable_thinking else "rgba(6, 182, 212, 0.15)"
active_mode_text = "#c4b5fd" if enable_thinking else "#a5f3fc"

st.markdown(
    f"""
    <div style="display: flex; justify-content: space-between; align-items: center; background: var(--bg-card); border: 1px solid var(--border); padding: 0.8rem 1.5rem; border-radius: 12px; margin-bottom: 2.2rem; box-shadow: 0 4px 20px rgba(0,0,0,0.25);">
        <div style="display: flex; align-items: center; gap: 8px;">
            <span style="height: 10px; width: 10px; background-color: #10b981; border-radius: 50%; display: inline-block; box-shadow: 0 0 8px #10b981;"></span>
            <span style="font-size: 0.85rem; font-weight: 600; color: var(--text-primary);">Server Online</span>
        </div>
        <div style="display: flex; gap: 16px; align-items: center; flex-wrap: wrap;">
            <span style="font-size: 0.85rem; color: var(--text-secondary);">Model: <code style="color: var(--accent-2); background: rgba(6, 182, 212, 0.05); padding: 2px 6px; border-radius: 4px; border: 1px solid rgba(6, 182, 212, 0.15);">{model_name or "Auto-detecting..."}</code></span>
            <span style="font-size: 0.85rem; padding: 4px 12px; border-radius: 20px; background: {active_mode_bg}; border: 1px solid {active_mode_color}; color: {active_mode_text}; font-weight: 600; display: inline-block;">{active_mode}</span>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# ── Tabs ─────────────────────────────────────────────────────────────────────

tab_image, tab_video, tab_docs, tab_text = st.tabs(
    ["🖼 Image", "🎬 Video", "📄 Documents (PDF/Docs)", "💬 Text-Only"]
)

# ── Image Tab ────────────────────────────────────────────────────────────────

with tab_image:
    img_source = st.radio(
        "Select Image Source",
        ["📁 Upload Image File(s)", "🔗 Image URL(s)", "📸 Camera Input"],
        horizontal=True,
        key="img_source",
    )

    grid_cols = st.slider("Preview Grid Columns", 1, 8, 4, key="img_grid_cols")

    uploaded_images = None
    img_urls_input = ""
    camera_image = None
    pil_images: list[Image.Image] = []

    if img_source == "📁 Upload Image File(s)":
        st.markdown(
            '<div class="section-label">Upload Images</div>',
            unsafe_allow_html=True,
        )
        uploaded_images = base64_uploader(
            "Drop images here",
            accept=["png", "jpg", "jpeg", "webp", "bmp", "gif"],
            multiple=True,
            key="img_upload",
        )
        st.info("💡 **Tip:** You can also paste an image from your clipboard by clicking the upload area and pressing **Ctrl+V**.")
        if uploaded_images:
            for img, fname in uploaded_files_to_images(uploaded_images):
                pil_images.append(img)

    elif img_source == "🔗 Image URL(s)":
        st.markdown(
            '<div class="section-label">Image URL(s)</div>',
            unsafe_allow_html=True,
        )
        img_urls_input = st.text_area(
            "Image URLs",
            placeholder="Enter image URLs (one per line or separated by commas)...",
            height=100,
            key="img_urls",
            label_visibility="collapsed",
        )
        if img_urls_input.strip():
            urls = [u.strip() for u in re.split(r'[\n,]+', img_urls_input) if u.strip()]
            for url in urls:
                if url.startswith(("http://", "https://")):
                    try:
                        r = requests.get(url, timeout=15)
                        if r.status_code == 200:
                            pil_images.append(Image.open(io.BytesIO(r.content)).convert("RGB"))
                        else:
                            st.error(f"❌ Failed to fetch image (HTTP {r.status_code}): {url}")
                    except Exception as e:
                        st.error(f"❌ Error loading image from URL ({url}): {e}")
                else:
                    st.warning(f"⚠️ Invalid URL (must start with http/https): {url}")

    else:
        st.markdown(
            '<div class="section-label">Camera Input</div>',
            unsafe_allow_html=True,
        )
        camera_image = st.camera_input("Take a photo", key="camera_input")
        if camera_image:
            try:
                pil_images.append(Image.open(camera_image).convert("RGB"))
            except Exception as e:
                st.error(f"❌ Failed to load camera image: {e}")

    if pil_images:
        st.markdown(
            f'<div class="section-label">📸 Input Images ({len(pil_images)})</div>',
            unsafe_allow_html=True,
        )
        cols = st.columns(min(len(pil_images), grid_cols))
        for i, img in enumerate(pil_images):
            with cols[i % len(cols)]:
                st.image(img, caption=f"Image {i+1}", use_container_width=True)

    img_prompt = st.text_area(
        "Image prompt",
        placeholder="Describe what you want to know about the image(s)…",
        height=100,
        key="img_prompt",
        label_visibility="collapsed",
    )

    if st.button("🔍 Analyze Image(s)", key="img_run", use_container_width=True):
        if not pil_images:
            st.warning("⚠ Please provide at least one image.")
        elif not img_prompt.strip():
            st.warning("⚠ Please enter a prompt.")
        else:
            messages = build_image_messages(img_prompt, pil_images, max_side=img_max_side)
            if use_streaming:
                render_stream(messages, max_tokens, temperature, model=model_name, enable_thinking=enable_thinking)
            else:
                render_non_stream(messages, max_tokens, temperature, model=model_name, enable_thinking=enable_thinking)

# ── Video Tab ────────────────────────────────────────────────────────────────

with tab_video:
    video_source = st.radio(
        "Select Video Source",
        ["📁 Upload Video File", "🔗 Video URL (YouTube, Vimeo, direct link, etc.)"],
        horizontal=True,
        key="vid_source",
    )

    video_grid_cols = st.slider("Frame Preview Grid Columns", 1, 8, 4, key="vid_grid_cols")

    uploaded_video_data = None
    youtube_url = ""

    if video_source == "📁 Upload Video File":
        st.markdown(
            '<div class="section-label">Upload Video</div>',
            unsafe_allow_html=True,
        )
        uploaded_video_data = base64_uploader(
            "Drop a video here",
            accept=["mp4", "avi", "mov", "mkv", "webm"],
            multiple=False,
            key="vid_upload",
        )
        if uploaded_video_data:
            vid_info = uploaded_video_data[0] if isinstance(uploaded_video_data, list) else uploaded_video_data
            if vid_info["size"] > 200 * 1024 * 1024:
                st.error("❌ Video must be under 200 MB for browser upload. Use a URL for larger files.")
                uploaded_video_data = None
            else:
                st.video(uploaded_file_to_bytes(vid_info))
    else:
        st.markdown(
            '<div class="section-label">Video URL</div>',
            unsafe_allow_html=True,
        )
        youtube_url = st.text_input(
            "Paste Video URL",
            placeholder="e.g. https://www.youtube.com/watch?v=dQw4w9WgXcQ",
            key="yt_url",
            label_visibility="collapsed",
        )
        if youtube_url.strip():
            if youtube_url.startswith(("http://", "https://")):
                try:
                    st.video(youtube_url)
                except Exception as e:
                    st.error(f"❌ Failed to load video: {e}")
            else:
                st.warning("⚠ Please enter a valid video URL (starting with http:// or https://).")

    vid_prompt = st.text_area(
        "Video prompt",
        placeholder="What do you want to know about this video?",
        height=100,
        key="vid_prompt",
        label_visibility="collapsed",
    )

    if st.button("🎬 Analyze Video", key="vid_run", use_container_width=True):
        if video_source == "📁 Upload Video File" and not uploaded_video_data:
            st.warning("⚠ Please upload a video.")
        elif video_source == "🔗 Video URL (YouTube, Vimeo, direct link, etc.)" and not youtube_url.strip():
            st.warning("⚠ Please enter a video URL.")
        elif not vid_prompt.strip():
            st.warning("⚠ Please enter a prompt.")
        else:
            frames = []
            tmp_path = None
            thumb_frames = None
            
            if video_source == "📁 Upload Video File":
                with st.spinner("Preparing uploaded video…"):
                    vid_info = uploaded_video_data[0] if isinstance(uploaded_video_data, list) else uploaded_video_data
                    vid_bytes = uploaded_file_to_bytes(vid_info)
                    with tempfile.NamedTemporaryFile(
                        delete=False, suffix=".mp4"
                    ) as tmp:
                        tmp.write(vid_bytes)
                        tmp_path = tmp.name
            else:
                with st.spinner("Downloading video (lowest/requested resolution)…"):
                    tmp_path, thumb_frames = download_youtube_video(youtube_url, quality=video_quality)

            if tmp_path:
                with st.spinner(f"Extracting up to {max_frames} frames…"):
                    frames = extract_frames(tmp_path, max_frames)
                
                # Clean up temporary video file
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
            elif thumb_frames:
                # Thumbnail fallback — use whatever frames we got from the CDN
                frames = thumb_frames
                st.info(f"📸 Using {len(frames)} thumbnail frames from YouTube CDN (limited coverage).")

            if not frames:
                st.error("❌ Could not extract any frames from the video.")
            else:
                st.markdown(
                    f'<div class="section-label">📸 Extracted {len(frames)} Frames</div>',
                    unsafe_allow_html=True,
                )
                frame_cols = st.columns(min(len(frames), video_grid_cols))
                for i, frame in enumerate(frames):
                    with frame_cols[i % len(frame_cols)]:
                        st.image(image_to_base64_url(frame), caption=f"Frame {i+1}", use_container_width=True)

                messages = build_video_messages(vid_prompt, frames, max_side=img_max_side)
                if use_streaming:
                    render_stream(messages, max_tokens, temperature, model=model_name, enable_thinking=enable_thinking)
                else:
                    render_non_stream(messages, max_tokens, temperature, model=model_name, enable_thinking=enable_thinking)

# ── Documents Tab ────────────────────────────────────────────────────────────

with tab_docs:
    st.markdown(
        '<div class="section-label">Upload Documents</div>',
        unsafe_allow_html=True,
    )
    uploaded_docs = base64_uploader(
        "Drop PDF or DOCX files here",
        accept=["pdf", "docx"],
        multiple=True,
        key="doc_upload",
    )

    parsed_docs = []
    if uploaded_docs:
        # We can extract text from each uploaded file and save the result
        with st.spinner("Extracting text from document(s)..."):
            for doc_info in uploaded_docs:
                filename = doc_info.get("name", "unknown")
                size_bytes = doc_info.get("size", 0)
                extracted_text, page_count = extract_text_from_file(doc_info)
                
                size_str = (
                    f"{size_bytes / 1024:.1f} KB"
                    if size_bytes < 1024 * 1024
                    else f"{size_bytes / (1024 * 1024):.1f} MB"
                )
                
                parsed_docs.append({
                    "name": filename,
                    "size_str": size_str,
                    "text": extracted_text,
                    "pages": page_count,
                    "chars": len(extracted_text),
                })
        
        # Display document details in a table-like layout
        st.markdown(
            '<div class="section-label">📄 Loaded Documents</div>',
            unsafe_allow_html=True,
        )
        for doc in parsed_docs:
            st.markdown(
                f"""
                <div style="background: rgba(255, 255, 255, 0.02); border: 1px solid var(--border); border-radius: 8px; padding: 0.6rem 1rem; margin-bottom: 0.5rem; display: flex; justify-content: space-between; align-items: center;">
                    <div>
                        <strong style="color: var(--accent-2);">{doc['name']}</strong>
                        <span style="font-size: 0.8rem; color: var(--text-secondary); margin-left: 10px;">({doc['size_str']})</span>
                    </div>
                    <div style="font-size: 0.8rem; color: var(--text-secondary);">
                        <span>Pages/Est. Pages: <strong>{doc['pages']}</strong></span>
                        <span style="margin: 0 10px;">|</span>
                        <span>Characters: <strong>{doc['chars']}</strong></span>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            
            with st.expander(f"👁 View Extracted Text: {doc['name']}", expanded=False):
                st.code(doc['text'][:5000] + ("\n... [TRUNCATED FOR PREVIEW] ..." if len(doc['text']) > 5000 else ""), language=None)

    doc_prompt = st.text_area(
        "Document prompt",
        placeholder="Ask Alpie anything about the uploaded document(s)…",
        height=100,
        key="doc_prompt",
        label_visibility="collapsed",
    )

    if st.button("🔍 Analyze Document(s)", key="doc_run", use_container_width=True):
        if not parsed_docs:
            st.warning("⚠ Please upload at least one document.")
        elif not doc_prompt.strip():
            st.warning("⚠ Please enter a prompt.")
        else:
            messages = build_document_messages(doc_prompt, parsed_docs)
            if use_streaming:
                render_stream(messages, max_tokens, temperature, model=model_name, enable_thinking=enable_thinking)
            else:
                render_non_stream(messages, max_tokens, temperature, model=model_name, enable_thinking=enable_thinking)

# ── Text Tab ─────────────────────────────────────────────────────────────────


with tab_text:
    st.markdown(
        '<div class="section-label">Text Prompt</div>',
        unsafe_allow_html=True,
    )

    text_prompt = st.text_area(
        "Text prompt",
        placeholder="Ask Alpie anything (text-only, no vision)…",
        height=130,
        key="text_prompt",
        label_visibility="collapsed",
    )

    if st.button("🚀 Run", key="text_run", use_container_width=True):
        if not text_prompt.strip():
            st.warning("⚠ Please enter a prompt.")
        else:
            messages = [{"role": "user", "content": text_prompt}]
            if use_streaming:
                render_stream(messages, max_tokens, temperature, model=model_name, enable_thinking=enable_thinking)
            else:
                render_non_stream(messages, max_tokens, temperature, model=model_name, enable_thinking=enable_thinking)

# ── Footer ───────────────────────────────────────────────────────────────────

st.markdown("---")
st.markdown(
    '<p style="text-align:center; color: var(--text-secondary); font-size:.8rem;">'
    "Alpie Vision Tester · Powered by 169Pi · "
    f'Endpoint: <code>{API_URL}</code></p>',
    unsafe_allow_html=True,
)
